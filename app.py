import io
import requests
import pandas as pd
import streamlit as st
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

import importlib.util, sys, subprocess
def _ensure(pkg_name, pip_name=None):
    if importlib.util.find_spec(pkg_name) is None:
        subprocess.check_call([sys.executable, "-m", "pip", "install",
                               pip_name or pkg_name, "lxml==5.3.0"])
_ensure("docx", "python-docx==1.1.2")
# --------------------------------------------------------

from docx import Document

st.set_page_config(page_title="LA Gran DT — Equivalencias", page_icon="📄", layout="wide")

st.title("📄 LA Gran DT — Generador de documento de equivalencias")
st.caption("Desde un Excel en OneDrive. Seleccioná filas y descargá el DOCX.")

with st.expander("Instrucciones", expanded=False):
    st.markdown("""
1. Compartí el Excel de OneDrive con enlace de lectura.
2. Pegá el enlace abajo.
3. Elegí hoja y asigná columnas.
4. Tildá las filas que quieras incluir.
5. Generá el Word.
    """)

def normalize_onedrive_url(url: str) -> str:
    if "download=1" not in url:
        sep = "&" if "?" in url else "?"
        url = f"{url}{sep}download=1"
    return url

@st.cache_data(show_spinner=False)
def fetch_excel(url: str) -> bytes:
    r = requests.get(normalize_onedrive_url(url), allow_redirects=True, timeout=60)
    r.raise_for_status()
    return r.content

def best_guess(cols, *cands):
    cols_low = [c.lower().strip() for c in cols]
    for group in cands:
        for cand in (group if isinstance(group, (list, tuple)) else [group]):
            for i, c in enumerate(cols_low):
                if cand in c:
                    return cols[i]
    return None

def build_docx(df: pd.DataFrame, alumno=None, legajo=None, carrera=None, periodo=None) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    title = doc.add_paragraph("Propuesta de Equivalencias ITBA ↔ POLIMI")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    if alumno:
        info.add_run(f"Alumno: {alumno}  ").bold = True
    if legajo:
        info.add_run(f"Legajo: {legajo}  ").bold = True
    if carrera:
        info.add_run(f"Carrera: {carrera}  ").bold = True
    if periodo:
        info.add_run(f"Período: {periodo}").bold = True

    doc.add_paragraph("")

    headers = ["Código ITBA", "Materia ITBA", "Créditos ITBA", "Código POLIMI", "Materia POLIMI", "ECTS"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row.get("itba_code", ""))
        cells[1].text = str(row.get("itba_name", ""))
        cells[2].text = str(row.get("itba_credits", ""))
        cells[3].text = str(row.get("polimi_code", ""))
        cells[4].text = str(row.get("polimi_name", ""))
        cells[5].text = str(row.get("polimi_ects", ""))

    doc.add_paragraph("")
    doc.add_paragraph("Generado automáticamente.")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

st.subheader("1) Enlace del Excel")
excel_url = st.text_input("Enlace público de OneDrive", placeholder="https://1drv.ms/x/s!...")

if excel_url:
    try:
        content = fetch_excel(excel_url)
        xls = pd.ExcelFile(BytesIO(content))
        st.success("Excel cargado")
        sheet = st.selectbox("Hoja", xls.sheet_names, index=0)
        df_raw = xls.parse(sheet)
        cols = list(df_raw.columns)

        st.subheader("2) Mapear columnas")
        guess_select = best_guess(cols, ["selección", "seleccion", "checkbox", "check", "créditos seleccionados", "creditos seleccionados"])
        guess_itba_code = best_guess(cols, ["itba"], ["cód", "cod", "codigo", "código"])
        guess_itba_name = best_guess(cols, ["materia itba", "materia", "itba"])
        guess_itba_credits = best_guess(cols, ["créditos itba", "creditos itba", "cr itba"])
        guess_polimi_code = best_guess(cols, ["polimi"], ["cód", "cod", "codigo", "código"], ["code"])
        guess_polimi_name = best_guess(cols, ["description", "descripción", "descripcion", "materia polimi", "polimi"])
        guess_polimi_ects = best_guess(cols, ["ects", "total", "créditos", "creditos"])

        c1, c2 = st.columns(2)
        with c1:
            col_select = st.selectbox("Columna de selección (opcional)", [None] + cols, index=(cols.index(guess_select)+1) if guess_select in cols else 0)
            itba_code = st.selectbox("Código ITBA", cols, index=cols.index(guess_itba_code) if guess_itba_code in cols else 0)
            itba_name = st.selectbox("Materia ITBA", cols, index=cols.index(guess_itba_name) if guess_itba_name in cols else 0)
            itba_credits = st.selectbox("Créditos ITBA", [None] + cols, index=(cols.index(guess_itba_credits)+1) if guess_itba_credits in cols else 0)
        with c2:
            polimi_code = st.selectbox("Código POLIMI", cols, index=cols.index(guess_polimi_code) if guess_polimi_code in cols else 0)
            polimi_name = st.selectbox("Materia POLIMI", cols, index=cols.index(guess_polimi_name) if guess_polimi_name in cols else 0)
            polimi_ects = st.selectbox("ECTS POLIMI", [None] + cols, index=(cols.index(guess_polimi_ects)+1) if guess_polimi_ects in cols else 0)

        work = pd.DataFrame({
            "itba_code": df_raw[itba_code],
            "itba_name": df_raw[itba_name],
            "itba_credits": df_raw[itba_credits] if itba_credits else "",
            "polimi_code": df_raw[polimi_code],
            "polimi_name": df_raw[polimi_name],
            "polimi_ects": df_raw[polimi_ects] if polimi_ects else "",
        })

        if col_select:
            sel = df_raw[col_select]
            work["incluir"] = sel.apply(lambda x: bool(x) and str(x).strip().lower() not in ["0", "no", "false", "f"])
        else:
            work["incluir"] = False

        st.subheader("3) Seleccionar filas")
        edited = st.data_editor(
            work,
            use_container_width=True,
            hide_index=True,
            column_config={"incluir": st.column_config.CheckboxColumn("Incluir", default=False)}
        )

        selected = edited[edited["incluir"]].drop(columns=["incluir"])
        st.write("Seleccionadas:", len(selected))

        st.subheader("4) Encabezado del documento")
        c3, c4, c5, c6 = st.columns(4)
        with c3:
            alumno = st.text_input("Alumno")
        with c4:
            legajo = st.text_input("Legajo")
        with c5:
            carrera = st.text_input("Carrera")
        with c6:
            periodo = st.text_input("Período")

        if st.button("Generar Word", type="primary", disabled=selected.empty):
            docx = build_docx(selected, alumno=alumno, legajo=legajo, carrera=carrera, periodo=periodo)
            st.download_button("Descargar DOCX", data=docx, file_name="LA_Gran_DT_equivalencias.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            csv = selected.to_csv(index=False).encode("utf-8")
            st.download_button("Descargar CSV", data=csv, file_name="seleccion.csv", mime="text/csv")

    except Exception as e:
        st.error(f"No se pudo leer el Excel. Revisá el enlace. Detalle: {e}")
else:
    st.info("Pegá el enlace público de OneDrive para continuar.")
